"""FastMCP server for Dexter filesystem, document, Outlook, and status tools."""

from __future__ import annotations

import datetime
import json
import logging
import os
import re
import threading
import time
from collections import deque
from enum import Enum
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any, Callable

from fastmcp import FastMCP

from core.event_bus import EventBus
from core.health import HealthMonitor, get_global_health_monitor, set_global_health_monitor
from tools.executor import ToolExecutor
from tools.schema_registry import load_tool_schemas
from utils.config import get_config

logger = logging.getLogger("dexter_mcp_server")

mcp = FastMCP("dexter-mcp-server")


class MCPErrorCode(str, Enum):
    TOOL_NOT_FOUND = "TOOL_NOT_FOUND"
    VALIDATION_FAILED = "VALIDATION_FAILED"
    RATE_LIMITED = "RATE_LIMITED"
    EXECUTION_ERROR = "EXECUTION_ERROR"
    INTERNAL_ERROR = "INTERNAL_ERROR"

ALLOWED_ROOTS: list[str] = json.loads(os.environ.get("DEXTER_ALLOWED_ROOTS", "[]"))
_SERVER_RUNTIME_ACTIVE = False
_EVENT_BUS = EventBus()
_LOCAL_HEALTH_MONITOR = HealthMonitor("dexter-mcp-server")
set_global_health_monitor(_LOCAL_HEALTH_MONITOR)

_TOOL_EXECUTOR: ToolExecutor | None = None
_TOOL_SCHEMAS: dict[str, dict] = {}
_RATE_LIMIT_WINDOWS: dict[str, deque[float]] = {}
_RATE_LIMIT_LOCK = threading.Lock()
_HTTP_SERVER: ThreadingHTTPServer | None = None
_HTTP_THREAD: threading.Thread | None = None

_TOOL_NAME_MAP = {
    "read_file": "mcp_read_file",
    "write_file": "mcp_write_file",
    "list_directory": "mcp_list_directory",
    "search_files": "mcp_search_files",
    "create_directory": "mcp_create_directory",
    "read_word_doc": "mcp_read_word_doc",
    "read_excel": "mcp_read_excel",
    "read_pdf": "mcp_read_pdf",
    "list_emails": "mcp_list_emails",
    "create_email_draft": "mcp_create_email_draft",
    "list_calendar_events": "mcp_list_calendar_events",
}


def _validate_path(path: str) -> Path:
    """Resolve a path and verify it sits under one of the configured roots."""
    resolved = Path(path).resolve()

    if not ALLOWED_ROOTS:
        raise PermissionError("No allowed roots configured. Cannot access filesystem.")

    for root in ALLOWED_ROOTS:
        root_resolved = Path(root).resolve()
        try:
            resolved.relative_to(root_resolved)
            return resolved
        except ValueError:
            continue

    raise PermissionError(
        f"Path '{path}' is outside all allowed roots. Allowed: {ALLOWED_ROOTS}"
    )


def _sanitize_health_for_external(summary: dict) -> dict:
    """
    Remove any key whose name contains a sensitive keyword.
    Operates recursively on nested dicts.
    """
    SENSITIVE_KEYWORDS = {"key", "secret", "token", "password", "api", "auth", "credential"}
    if not isinstance(summary, dict):
        return summary
    cleaned = {}
    for k, v in summary.items():
        key_lower = k.lower()
        if any(kw in key_lower for kw in SENSITIVE_KEYWORDS):
            continue  # drop this field entirely
        if isinstance(v, dict):
            cleaned[k] = _sanitize_health_for_external(v)
        elif isinstance(v, list):
            cleaned[k] = [
                _sanitize_health_for_external(item) if isinstance(item, dict) else item
                for item in v
            ]
        else:
            cleaned[k] = v
    return cleaned


def _schema_name(tool_name: str) -> str:
    return _TOOL_NAME_MAP.get(tool_name, f"mcp_{tool_name}")


def _client_id() -> str:
    return (
        os.environ.get("DEXTER_MCP_CLIENT_ID")
        or os.environ.get("MCP_CLIENT_ID")
        or "stdio"
    )


def _rate_limit_per_minute() -> int:
    try:
        configured = int(get_config().mcp.max_calls_per_minute)
        return max(1, configured)
    except Exception:
        return 30


def _tool_error_response(
    tool_name: str,
    client_id: str,
    code: MCPErrorCode,
    message: str,
    *,
    details: dict[str, Any] | None = None,
) -> dict[str, Any]:
    payload: dict[str, Any] = {
        "success": False,
        "error": {
            "code": code,
            "message": message,
            "tool": tool_name,
            "client_id": client_id,
        },
    }
    if details:
        payload["error"]["details"] = details
    return payload


def _map_tool_error_code(code: str) -> MCPErrorCode:
    normalized = (code or "").strip().lower()
    if normalized == "rate_limited":
        return MCPErrorCode.RATE_LIMITED
    if normalized in {"validation_failed", "confirmation_required"}:
        return MCPErrorCode.VALIDATION_FAILED
    if normalized in {"tool_unavailable", "tool_not_found", "schema_missing"}:
        return MCPErrorCode.TOOL_NOT_FOUND
    if normalized in {"execution_error", "error"}:
        return MCPErrorCode.EXECUTION_ERROR
    return MCPErrorCode.INTERNAL_ERROR


def _emit_mcp_event(tool_name: str, client_id: str, status: str, **details: Any) -> None:
    _EVENT_BUS.emit(
        "mcp_tool_call",
        {
            "tool": tool_name,
            "client_id": client_id,
            "status": status,
            **details,
        },
    )


def _rate_limit_check(client_id: str) -> tuple[bool, float]:
    limit = _rate_limit_per_minute()
    now = time.time()
    with _RATE_LIMIT_LOCK:
        window = _RATE_LIMIT_WINDOWS.setdefault(client_id, deque())
        while window and now - window[0] >= 60.0:
            window.popleft()
        if len(window) >= limit:
            retry_after = 60.0 - (now - window[0]) if window else 60.0
            return False, max(0.0, retry_after)
        window.append(now)
        return True, 0.0


def _build_proxy(tool_name: str, impl: Callable[..., Any]) -> Callable[..., Any]:
    async def _proxy(**kwargs: Any) -> Any:
        return await impl(**kwargs)

    _proxy.__name__ = tool_name
    _proxy.__doc__ = impl.__doc__
    return _proxy


def _get_tool_executor() -> ToolExecutor:
    global _TOOL_EXECUTOR, _TOOL_SCHEMAS
    if _TOOL_EXECUTOR is None:
        proxies = [_build_proxy(name, impl) for name, impl in _TOOL_IMPLEMENTATIONS.items()]
        executor = ToolExecutor(proxies, event_bus=_EVENT_BUS)
        schemas = load_tool_schemas()
        for tool_name in _TOOL_IMPLEMENTATIONS:
            schema = schemas.get(_schema_name(tool_name))
            if schema:
                executor._schemas[tool_name] = schema
                _TOOL_SCHEMAS[tool_name] = schema
        _TOOL_EXECUTOR = executor
    return _TOOL_EXECUTOR


async def _dispatch_tool(tool_name: str, args: dict[str, Any]) -> dict[str, Any]:
    client_id = _client_id()
    allowed, retry_after = _rate_limit_check(client_id)
    if not allowed:
        message = (
            f"Rate limit exceeded for client '{client_id}'. "
            f"Max {_rate_limit_per_minute()} calls per minute."
        )
        _emit_mcp_event(tool_name, client_id, "rate_limited", retry_after_seconds=round(retry_after, 1))
        return _tool_error_response(
            tool_name,
            client_id,
            MCPErrorCode.RATE_LIMITED,
            message,
            details={"retry_after_seconds": round(retry_after, 1)},
        )

    schema_key = _schema_name(tool_name)
    schema = load_tool_schemas().get(schema_key)
    if not schema:
        _emit_mcp_event(tool_name, client_id, "validation_failed", reason="schema_missing")
        return _tool_error_response(
            tool_name,
            client_id,
            MCPErrorCode.TOOL_NOT_FOUND,
            f"Schema '{schema_key}' is not registered.",
        )

    executor = _get_tool_executor()
    if tool_name not in executor._tools:
        _emit_mcp_event(tool_name, client_id, "validation_failed", reason="tool_unavailable")
        return _tool_error_response(
            tool_name,
            client_id,
            MCPErrorCode.TOOL_NOT_FOUND,
            f"Tool '{tool_name}' is not registered.",
        )

    result = await executor.execute(tool_name, args or {}, event_bus=_EVENT_BUS)
    if result.success:
        _emit_mcp_event(tool_name, client_id, "success", duration_ms=round(result.duration_ms, 1))
        return result.data

    error_text = result.error or "Tool call failed."
    error_code = MCPErrorCode.VALIDATION_FAILED
    if result.policy_decision == "tool_not_available":
        error_code = MCPErrorCode.TOOL_NOT_FOUND
    elif result.confirmation_required:
        error_code = MCPErrorCode.VALIDATION_FAILED
    elif result.policy_decision not in {"invalid_args", "tool_not_available"}:
        error_code = _map_tool_error_code(result.policy_decision or "validation_failed")

    _emit_mcp_event(
        tool_name,
        client_id,
        "validation_failed",
        duration_ms=round(result.duration_ms, 1),
        policy_decision=result.policy_decision,
    )
    return _tool_error_response(
        tool_name,
        client_id,
        error_code,
        error_text,
        details={
            "duration_ms": round(result.duration_ms, 1),
            "risk_level": result.risk_level,
            "confirmation_required": result.confirmation_required,
            "policy_decision": result.policy_decision,
        },
    )


def _impl_read_file(path: str) -> dict:
    validated = _validate_path(path)

    size = validated.stat().st_size
    if size > 10 * 1024 * 1024:
        return {"success": False, "error": f"File too large: {size} bytes. Maximum is 10MB."}

    content = None
    encoding_used = None
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            content = validated.read_text(encoding=enc)
            encoding_used = enc
            break
        except UnicodeDecodeError:
            continue

    if content is None:
        return {"success": False, "error": "File appears to be binary. Cannot read as text."}

    return {
        "success": True,
        "content": content,
        "encoding": encoding_used,
        "size_bytes": size,
        "lines": content.count("\n") + 1,
        "path": str(validated),
    }


def _impl_write_file(path: str, content: str, mode: str = "overwrite") -> dict:
    if mode not in ("overwrite", "append"):
        return {"success": False, "error": f"Invalid mode '{mode}'. Must be overwrite or append."}

    validated = _validate_path(path)
    validated.parent.mkdir(parents=True, exist_ok=True)

    if mode == "overwrite":
        validated.write_text(content, encoding="utf-8")
    else:
        with open(validated, "a", encoding="utf-8") as handle:
            handle.write(content)

    bytes_written = len(content.encode("utf-8"))

    return {"success": True, "bytes_written": bytes_written, "path": str(validated), "mode": mode}


def _impl_list_directory(path: str, pattern: str = "*", include_hidden: bool = False) -> dict:
    validated = _validate_path(path)

    if not validated.is_dir():
        return {"success": False, "error": f"Not a directory: {path}"}

    items = []
    for item in validated.glob(pattern):
        if not include_hidden and item.name.startswith("."):
            continue

        try:
            stat = item.stat()
            items.append(
                {
                    "name": item.name,
                    "type": "directory" if item.is_dir() else "file",
                    "size_bytes": stat.st_size if item.is_file() else 0,
                    "modified_iso": datetime.datetime.fromtimestamp(stat.st_mtime).isoformat(),
                    "extension": item.suffix.lower() if item.is_file() else "",
                }
            )
        except (PermissionError, OSError):
            continue

        if len(items) >= 500:
            return {
                "success": True,
                "items": items,
                "total_count": len(items),
                "truncated": True,
                "message": "Results truncated at 500. Use pattern to filter.",
            }

    return {
        "success": True,
        "items": sorted(items, key=lambda x: (x["type"] == "file", x["name"].lower())),
        "total_count": len(items),
        "truncated": False,
    }


def _impl_search_files(
    root: str,
    query: str,
    file_types: list | None = None,
    search_content: bool = False,
    max_results: int = 50,
) -> dict:
    start = time.time()
    deadline = start + 15.0

    validated_root = _validate_path(root)
    max_results = min(max_results, 100)

    query_lower = query.lower()
    matches = []

    extensions = None
    if file_types:
        extensions = {f".{ext.lstrip('.')}" for ext in file_types}

    for file_path in validated_root.rglob("*"):
        if time.time() > deadline:
            break

        if not file_path.is_file():
            continue

        if extensions and file_path.suffix.lower() not in extensions:
            continue

        if any(part.startswith(".") for part in file_path.parts):
            continue

        if any(
            d in file_path.parts
            for d in (
                "node_modules",
                "__pycache__",
                ".venv",
                "venv",
                ".git",
                "build",
                "dist",
                "target",
            )
        ):
            continue

        if query_lower in file_path.name.lower():
            matches.append(
                {
                    "path": str(file_path),
                    "name": file_path.name,
                    "match_type": "filename",
                    "snippet": "",
                }
            )
            if len(matches) >= max_results:
                break
            continue

        if search_content and file_path.stat().st_size < 1024 * 1024:
            try:
                text = file_path.read_text(encoding="utf-8", errors="ignore")
                idx = text.lower().find(query_lower)
                if idx >= 0:
                    start_idx = max(0, idx - 100)
                    end_idx = min(len(text), idx + 200)
                    snippet = text[start_idx:end_idx]
                    snippet = snippet.replace("\n", " ").strip()

                    matches.append(
                        {
                            "path": str(file_path),
                            "name": file_path.name,
                            "match_type": "content",
                            "snippet": snippet,
                        }
                    )

                    if len(matches) >= max_results:
                        break
            except Exception:
                continue

    elapsed = (time.time() - start) * 1000

    return {
        "success": True,
        "matches": matches,
        "total_matches": len(matches),
        "search_time_ms": round(elapsed, 1),
        "truncated": len(matches) >= max_results,
    }


def _impl_create_directory(path: str) -> dict:
    validated = _validate_path(path)
    already_existed = validated.exists()
    validated.mkdir(parents=True, exist_ok=True)
    return {"success": True, "path": str(validated), "already_existed": already_existed}


def _impl_read_word_doc(path: str) -> dict:
    validated = _validate_path(path)

    if not str(validated).lower().endswith(".docx"):
        return {"success": False, "error": "Only .docx files supported. Not .doc or other formats."}

    if validated.name.startswith("~$"):
        return {"success": False, "error": "Temp file (still open in Word). Close Word first."}

    from docx import Document

    doc = Document(str(validated))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)

    full_text = "\n".join(paragraphs)
    word_count = len(full_text.split())

    title = ""
    try:
        title = doc.core_properties.title or ""
    except Exception:
        pass

    return {
        "success": True,
        "title": title,
        "paragraphs": paragraphs,
        "tables": tables,
        "full_text": full_text,
        "word_count": word_count,
        "page_estimate": max(1, word_count // 250),
    }


def _impl_read_excel(path: str, sheet: str | None = None, max_rows: int = 1000) -> dict:
    validated = _validate_path(path)
    max_rows = min(max_rows, 5000)

    import openpyxl

    wb = openpyxl.load_workbook(str(validated), read_only=True, data_only=True)

    all_sheets = wb.sheetnames

    if sheet:
        if sheet not in all_sheets:
            wb.close()
            return {"success": False, "error": f"Sheet '{sheet}' not found. Available: {all_sheets}"}
        ws = wb[sheet]
    else:
        ws = wb.active

    sheet_name = ws.title
    rows = []
    headers = []

    for i, row in enumerate(ws.iter_rows(values_only=True)):
        row_data = ["" if v is None else str(v) for v in row]

        if i == 0:
            headers = row_data
        else:
            rows.append(row_data)

        if len(rows) >= max_rows:
            wb.close()
            return {
                "success": True,
                "sheet_name": sheet_name,
                "all_sheet_names": all_sheets,
                "headers": headers,
                "rows": rows,
                "row_count": len(rows),
                "column_count": len(headers),
                "truncated": True,
            }

    wb.close()

    return {
        "success": True,
        "sheet_name": sheet_name,
        "all_sheet_names": all_sheets,
        "headers": headers,
        "rows": rows,
        "row_count": len(rows),
        "column_count": len(headers),
        "truncated": False,
    }


def _impl_read_pdf(path: str, max_pages: int = 50) -> dict:
    validated = _validate_path(path)
    max_pages = min(max_pages, 200)

    import pypdf

    with open(str(validated), "rb") as handle:
        reader = pypdf.PdfReader(handle)

        if reader.is_encrypted:
            return {"success": False, "error": "PDF is encrypted. Cannot read without password."}

        total_pages = len(reader.pages)
        pages_to_read = min(total_pages, max_pages)

        title = ""
        author = ""
        try:
            meta = reader.metadata
            if meta:
                title = meta.get("/Title", "") or ""
                author = meta.get("/Author", "") or ""
        except Exception:
            pass

        pages = []
        full_text_parts = []

        for i in range(pages_to_read):
            try:
                text = reader.pages[i].extract_text()
                if not text or not text.strip():
                    text = "[scanned page - no text]"
            except Exception:
                text = "[page extraction failed]"

            pages.append({"page_number": i + 1, "text": text})
            full_text_parts.append(f"--- Page {i + 1} ---\n{text}")

        return {
            "success": True,
            "title": title,
            "author": author,
            "page_count": total_pages,
            "pages_extracted": pages_to_read,
            "pages": pages,
            "full_text": "\n\n".join(full_text_parts),
            "truncated": total_pages > max_pages,
        }


def _impl_list_emails(folder: str = "Inbox", count: int = 10, unread_only: bool = False) -> dict:
    import pythoncom

    pythoncom.CoInitialize()

    try:
        import win32com.client

        outlook = win32com.client.Dispatch("Outlook.Application")
        namespace = outlook.GetNamespace("MAPI")

        try:
            inbox = namespace.GetDefaultFolder(6)
            if folder == "Inbox":
                target = inbox
            elif folder == "Sent Items":
                target = namespace.GetDefaultFolder(5)
            elif folder == "Drafts":
                target = namespace.GetDefaultFolder(16)
            else:
                target = inbox.Parent.Folders[folder]
        except Exception:
            return {"success": False, "error": f"Folder '{folder}' not found."}

        messages = target.Items
        messages.Sort("[ReceivedTime]", True)

        emails = []
        count = min(count, 50)

        for msg in messages:
            if len(emails) >= count:
                break

            try:
                if unread_only and msg.UnRead is False:
                    continue

                body = msg.Body or ""
                body_clean = re.sub(r"<[^>]+>", "", body)[:200].strip()

                emails.append(
                    {
                        "subject": msg.Subject or "",
                        "sender_name": msg.SenderName or "",
                        "sender_email": msg.SenderEmailAddress or "",
                        "received_time": str(msg.ReceivedTime),
                        "is_unread": bool(msg.UnRead),
                        "body_preview": body_clean,
                        "has_attachments": msg.Attachments.Count > 0,
                    }
                )
            except Exception:
                continue

        return {"success": True, "folder": folder, "emails": emails, "total_returned": len(emails)}

    finally:
        pythoncom.CoUninitialize()


def _impl_create_email_draft(to: str, subject: str, body: str, cc: str = "") -> dict:
    import pythoncom

    pythoncom.CoInitialize()

    try:
        import win32com.client

        outlook = win32com.client.Dispatch("Outlook.Application")

        mail = outlook.CreateItem(0)
        mail.To = to
        mail.Subject = subject
        mail.Body = body
        if cc:
            mail.CC = cc

        mail.Display(True)

        return {
            "success": True,
            "action": "draft_created_and_displayed",
            "to": to,
            "subject": subject,
            "message": "Draft opened in Outlook. Review and send manually.",
        }

    finally:
        pythoncom.CoUninitialize()


def _impl_list_calendar_events(days_ahead: int = 7, days_back: int = 0) -> dict:
    import datetime as dt
    import pythoncom

    pythoncom.CoInitialize()

    try:
        import win32com.client

        outlook = win32com.client.Dispatch("Outlook.Application")
        namespace = outlook.GetNamespace("MAPI")
        calendar = namespace.GetDefaultFolder(9)

        days_ahead = min(days_ahead, 90)
        days_back = min(days_back, 30)

        now = dt.datetime.now()
        start_date = now - dt.timedelta(days=days_back)
        end_date = now + dt.timedelta(days=days_ahead)

        items = calendar.Items
        items.IncludeRecurrences = True
        items.Sort("[Start]")

        restriction = (
            f"[Start] >= '{start_date.strftime('%m/%d/%Y')}' "
            f"AND [Start] <= '{end_date.strftime('%m/%d/%Y')}'"
        )

        filtered = items.Restrict(restriction)
        events = []

        for event in filtered:
            try:
                events.append(
                    {
                        "subject": event.Subject or "",
                        "start": str(event.Start),
                        "end": str(event.End),
                        "location": event.Location or "",
                        "organizer": event.Organizer or "",
                        "is_all_day": bool(event.AllDayEvent),
                        "body_preview": (event.Body or "")[:150].strip(),
                    }
                )
            except Exception:
                continue

        return {
            "success": True,
            "events": sorted(events, key=lambda e: e["start"]),
            "total_events": len(events),
            "range_start": start_date.date().isoformat(),
            "range_end": end_date.date().isoformat(),
        }

    finally:
        pythoncom.CoUninitialize()


_TOOL_IMPLEMENTATIONS: dict[str, Callable[..., Any]] = {
    "read_file": _impl_read_file,
    "write_file": _impl_write_file,
    "list_directory": _impl_list_directory,
    "search_files": _impl_search_files,
    "create_directory": _impl_create_directory,
    "read_word_doc": _impl_read_word_doc,
    "read_excel": _impl_read_excel,
    "read_pdf": _impl_read_pdf,
    "list_emails": _impl_list_emails,
    "create_email_draft": _impl_create_email_draft,
    "list_calendar_events": _impl_list_calendar_events,
}


@mcp.tool()
async def read_file(path: str) -> dict:
    """Read the complete text content of a file within allowed roots (max 10MB)."""
    if not _SERVER_RUNTIME_ACTIVE:
        try:
            return _impl_read_file(path)
        except PermissionError as e:
            return {"success": False, "error": str(e)}
        except FileNotFoundError:
            return {"success": False, "error": f"File not found: {path}"}
        except Exception as e:
            return {"success": False, "error": f"Unexpected error: {str(e)}"}
    return await _dispatch_tool("read_file", {"path": path})


@mcp.tool()
async def write_file(path: str, content: str, mode: str = "overwrite") -> dict:
    """Write text content to a file (overwrite or append)."""
    if not _SERVER_RUNTIME_ACTIVE:
        try:
            return _impl_write_file(path, content, mode)
        except PermissionError as e:
            return {"success": False, "error": str(e)}
        except Exception as e:
            return {"success": False, "error": f"Write failed: {str(e)}"}
    return await _dispatch_tool("write_file", {"path": path, "content": content, "mode": mode})


@mcp.tool()
async def list_directory(path: str, pattern: str = "*", include_hidden: bool = False) -> dict:
    """List files and folders at the given path (up to 500 items)."""
    if not _SERVER_RUNTIME_ACTIVE:
        try:
            return _impl_list_directory(path, pattern, include_hidden)
        except PermissionError as e:
            return {"success": False, "error": str(e)}
        except Exception as e:
            return {"success": False, "error": f"List failed: {str(e)}"}
    return await _dispatch_tool(
        "list_directory",
        {"path": path, "pattern": pattern, "include_hidden": include_hidden},
    )


@mcp.tool()
async def search_files(
    root: str,
    query: str,
    file_types: list | None = None,
    search_content: bool = False,
    max_results: int = 50,
) -> dict:
    """Search for files by name and optionally content under root."""
    if not _SERVER_RUNTIME_ACTIVE:
        try:
            return _impl_search_files(root, query, file_types, search_content, max_results)
        except PermissionError as e:
            return {"success": False, "error": str(e)}
        except Exception as e:
            return {"success": False, "error": f"Search failed: {str(e)}"}
    return await _dispatch_tool(
        "search_files",
        {
            "root": root,
            "query": query,
            "file_types": file_types,
            "search_content": search_content,
            "max_results": max_results,
        },
    )


@mcp.tool()
async def create_directory(path: str) -> dict:
    """Create a directory and all parent directories."""
    if not _SERVER_RUNTIME_ACTIVE:
        try:
            return _impl_create_directory(path)
        except PermissionError as e:
            return {"success": False, "error": str(e)}
        except Exception as e:
            return {"success": False, "error": f"Create failed: {str(e)}"}
    return await _dispatch_tool("create_directory", {"path": path})


@mcp.tool()
async def read_word_doc(path: str) -> dict:
    """Extract text from a Microsoft Word .docx file."""
    if not _SERVER_RUNTIME_ACTIVE:
        try:
            return _impl_read_word_doc(path)
        except PermissionError as e:
            return {"success": False, "error": str(e)}
        except Exception as e:
            return {"success": False, "error": f"Could not read Word doc: {str(e)}"}
    return await _dispatch_tool("read_word_doc", {"path": path})


@mcp.tool()
async def read_excel(path: str, sheet: str | None = None, max_rows: int = 1000) -> dict:
    """Read data from a Microsoft Excel .xlsx file."""
    if not _SERVER_RUNTIME_ACTIVE:
        try:
            return _impl_read_excel(path, sheet, max_rows)
        except PermissionError as e:
            return {"success": False, "error": str(e)}
        except Exception as e:
            return {"success": False, "error": f"Could not read Excel: {str(e)}"}
    return await _dispatch_tool("read_excel", {"path": path, "sheet": sheet, "max_rows": max_rows})


@mcp.tool()
async def read_pdf(path: str, max_pages: int = 50) -> dict:
    """Extract text from a PDF file."""
    if not _SERVER_RUNTIME_ACTIVE:
        try:
            return _impl_read_pdf(path, max_pages)
        except PermissionError as e:
            return {"success": False, "error": str(e)}
        except Exception as e:
            return {"success": False, "error": f"Could not read PDF: {str(e)}"}
    return await _dispatch_tool("read_pdf", {"path": path, "max_pages": max_pages})


@mcp.tool()
async def list_emails(folder: str = "Inbox", count: int = 10, unread_only: bool = False) -> dict:
    """List recent emails from Outlook."""
    if not _SERVER_RUNTIME_ACTIVE:
        try:
            return _impl_list_emails(folder, count, unread_only)
        except ImportError:
            return {"success": False, "error": "pywin32 not installed or not running on Windows."}
        except Exception as e:
            return {"success": False, "error": f"Outlook error: {str(e)}"}
    return await _dispatch_tool(
        "list_emails",
        {"folder": folder, "count": count, "unread_only": unread_only},
    )


@mcp.tool()
async def create_email_draft(to: str, subject: str, body: str, cc: str = "") -> dict:
    """Create an email draft in Outlook (never auto-sends)."""
    if not _SERVER_RUNTIME_ACTIVE:
        try:
            return _impl_create_email_draft(to, subject, body, cc)
        except ImportError:
            return {"success": False, "error": "pywin32 not installed."}
        except Exception as e:
            return {"success": False, "error": f"Could not create draft: {str(e)}"}
    return await _dispatch_tool(
        "create_email_draft",
        {"to": to, "subject": subject, "body": body, "cc": cc},
    )


@mcp.tool()
async def list_calendar_events(days_ahead: int = 7, days_back: int = 0) -> dict:
    """List calendar events from Outlook Calendar."""
    if not _SERVER_RUNTIME_ACTIVE:
        try:
            return _impl_list_calendar_events(days_ahead, days_back)
        except ImportError:
            return {"success": False, "error": "pywin32 not installed."}
        except Exception as e:
            return {"success": False, "error": f"Calendar error: {str(e)}"}
    return await _dispatch_tool(
        "list_calendar_events",
        {"days_ahead": days_ahead, "days_back": days_back},
    )


class _MCPStatusHandler(BaseHTTPRequestHandler):
    server_version = "DexterMCP/1.0"

    @staticmethod
    def _redact_sensitive(value: Any) -> Any:
        if isinstance(value, dict):
            redacted: dict[str, Any] = {}
            for key, item in value.items():
                key_name = str(key).lower()
                if any(token in key_name for token in ("token", "secret", "password", "api_key", "key")):
                    redacted[key] = "[redacted]"
                else:
                    redacted[key] = _MCPStatusHandler._redact_sensitive(item)
            return redacted
        if isinstance(value, list):
            return [_MCPStatusHandler._redact_sensitive(item) for item in value]
        return value

    def _write_json(self, status_code: int, payload: dict[str, Any]) -> None:
        body = json.dumps(payload, ensure_ascii=False, default=str).encode("utf-8")
        self.send_response(status_code)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_GET(self) -> None:  # noqa: N802
        if self.path == "/health":
            monitor = get_global_health_monitor() or _LOCAL_HEALTH_MONITOR
            raw_summary = monitor.get_health_summary()
            safe_summary = _sanitize_health_for_external(raw_summary)
            self._write_json(200, safe_summary)
            return
        if self.path == "/tools":
            self._write_json(200, {"tools": load_tool_schemas()})
            return
        if self.path == "/":
            self._write_json(
                200,
                {
                    "service": "dexter-mcp-server",
                    "endpoints": ["/health", "/tools"],
                },
            )
            return
        self._write_json(404, {"success": False, "error": "Not found"})

    def log_message(self, format: str, *args: Any) -> None:  # noqa: A003
        logger.info("mcp_http_request", path=self.path, message=format % args if args else format)


def _start_http_sidecar() -> None:
    global _HTTP_SERVER, _HTTP_THREAD
    if _HTTP_SERVER is not None:
        return

    try:
        port = int(os.environ.get("DEXTER_MCP_HTTP_PORT", "8765"))
    except ValueError:
        port = 8765

    if port <= 0:
        logger.info("mcp_http_sidecar_disabled")
        return

    try:
        _HTTP_SERVER = ThreadingHTTPServer(("127.0.0.1", port), _MCPStatusHandler)
    except OSError as exc:
        logger.warning("mcp_http_sidecar_failed", port=port, error=str(exc))
        _HTTP_SERVER = None
        return

    def _serve() -> None:
        assert _HTTP_SERVER is not None
        logger.info("mcp_http_sidecar_started", host="127.0.0.1", port=port)
        _HTTP_SERVER.serve_forever(poll_interval=0.5)

    _HTTP_THREAD = threading.Thread(target=_serve, name="dexter-mcp-http", daemon=True)
    _HTTP_THREAD.start()


def main() -> None:
    global _SERVER_RUNTIME_ACTIVE
    _SERVER_RUNTIME_ACTIVE = True
    _get_tool_executor()
    _start_http_sidecar()
    mcp.run(transport="stdio")


if __name__ == "__main__":
    main()
